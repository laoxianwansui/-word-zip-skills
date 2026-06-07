"""解析 .docx 文件，提取段落文本、样式和格式信息。"""

from dataclasses import dataclass, field
from typing import Optional


@dataclass
class FormattedSpan:
    """一个格式化文本片段"""
    text: str
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    bold: bool = False
    italic: bool = False


@dataclass
class Paragraph:
    """解析后的段落"""
    index: int
    text: str
    spans: list[FormattedSpan] = field(default_factory=list)
    style_name: Optional[str] = None
    alignment: Optional[str] = None
    first_line_indent_pt: Optional[float] = None
    line_spacing: Optional[float] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    is_table_cell: bool = False
    page_break_before: bool = False

    def dominant_font_name(self) -> Optional[str]:
        if not self.spans:
            return None
        named_spans = [s for s in self.spans if s.font_name]
        if not named_spans:
            return None
        return max(named_spans, key=lambda s: len(s.text)).font_name

    def dominant_font_size_pt(self) -> Optional[float]:
        if not self.spans:
            return None
        sized_spans = [s for s in self.spans if s.font_size_pt is not None]
        if not sized_spans:
            return None
        return max(sized_spans, key=lambda s: len(s.text)).font_size_pt

    def is_bold(self) -> bool:
        if not self.spans:
            return False
        bold_chars = sum(len(s.text) for s in self.spans if s.bold)
        total_chars = sum(len(s.text) for s in self.spans)
        return bold_chars > total_chars * 0.5 if total_chars > 0 else False

    def is_empty(self) -> bool:
        return len(self.text.strip()) == 0


@dataclass
class TableInfo:
    """表格信息"""
    index: int
    rows: int
    cols: int
    headers: list[str] = field(default_factory=list)
    cells: list[list[str]] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """统一的文档解析结果"""
    file_path: str
    file_type: str
    paragraphs: list[Paragraph] = field(default_factory=list)
    tables: list[TableInfo] = field(default_factory=list)
    total_pages_estimate: int = 0
    raw_metadata: dict = field(default_factory=dict)
    parser_backend: str = "python-docx"
    parser_warnings: list[str] = field(default_factory=list)


def _emu_to_pt(emu: Optional[int]) -> Optional[float]:
    if emu is None:
        return None
    return emu / 12700.0


def _twips_to_pt(twips: Optional[int]) -> Optional[float]:
    if twips is None:
        return None
    return twips / 20.0


def _alignment_to_str(align) -> Optional[str]:
    if align is None:
        return None
    try:
        align = int(align)
    except Exception:
        pass
    align_map = {
        0: "left",
        1: "center",
        2: "right",
        3: "justify",
    }
    return align_map.get(align)


def _first_not_none(*values):
    for value in values:
        if value is not None:
            return value
    return None


def _style_chain(style):
    seen = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        yield current
        current = getattr(current, "base_style", None)


def _font_name_from_font(font) -> Optional[str]:
    if font is None:
        return None
    return getattr(font, "name", None) or None


def _font_size_from_font(font) -> Optional[float]:
    if font is None or getattr(font, "size", None) is None:
        return None
    return _emu_to_pt(font.size)


def _bold_from_font(font) -> Optional[bool]:
    if font is None:
        return None
    return getattr(font, "bold", None)


def _italic_from_font(font) -> Optional[bool]:
    if font is None:
        return None
    return getattr(font, "italic", None)


def _paragraph_format_value(paragraph_format, attr: str):
    if paragraph_format is None:
        return None
    return getattr(paragraph_format, attr, None)


def _length_to_pt(value) -> Optional[float]:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return value.pt
    try:
        return _twips_to_pt(value)
    except Exception:
        return None


def _east_asia_font(run) -> Optional[str]:
    rpr = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rpr is None:
        return None
    return rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')


def parse_docx(file_path: str) -> ParsedDocument:
    """解析 .docx 文件，返回统一的 ParsedDocument 对象。"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = Document(file_path)
    doc_info = ParsedDocument(
        file_path=file_path,
        file_type="docx",
        parser_backend="python-docx",
    )

    try:
        core_props = doc.core_properties
        doc_info.raw_metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
    except Exception:
        pass

    def _extract_spans(paragraph) -> list[FormattedSpan]:
        spans = []
        paragraph_styles = list(_style_chain(paragraph.style)) if paragraph.style else []

        for run in paragraph.runs:
            run_styles = list(_style_chain(run.style)) if run.style else []
            fonts = [run.font]
            fonts.extend(style.font for style in run_styles if getattr(style, "font", None) is not None)
            fonts.extend(style.font for style in paragraph_styles if getattr(style, "font", None) is not None)

            font_name = _east_asia_font(run) or _first_not_none(*[_font_name_from_font(font) for font in fonts])
            font_size = _first_not_none(*[_font_size_from_font(font) for font in fonts])
            bold = _first_not_none(*[_bold_from_font(font) for font in fonts], False)
            italic = _first_not_none(*[_italic_from_font(font) for font in fonts], False)

            spans.append(FormattedSpan(
                text=run.text,
                font_name=font_name,
                font_size_pt=font_size,
                bold=bool(bold),
                italic=bool(italic),
            ))

        if not spans and paragraph.text:
            fonts = [style.font for style in paragraph_styles if getattr(style, "font", None) is not None]
            spans.append(FormattedSpan(
                text=paragraph.text,
                font_name=_first_not_none(*[_font_name_from_font(font) for font in fonts]),
                font_size_pt=_first_not_none(*[_font_size_from_font(font) for font in fonts]),
                bold=bool(_first_not_none(*[_bold_from_font(font) for font in fonts], False)),
                italic=bool(_first_not_none(*[_italic_from_font(font) for font in fonts], False)),
            ))
        return spans

    def _extract_paragraph(paragraph, idx: int) -> Paragraph:
        para_format = paragraph.paragraph_format
        style_formats = [style.paragraph_format for style in _style_chain(paragraph.style)] if paragraph.style else []
        formats = [para_format] + style_formats

        first_indent = _length_to_pt(_first_not_none(*[
            _paragraph_format_value(fmt, "first_line_indent") for fmt in formats
        ]))

        line_spacing_value = _first_not_none(*[
            _paragraph_format_value(fmt, "line_spacing") for fmt in formats
        ])
        line_spacing = None
        if line_spacing_value is not None:
            if hasattr(line_spacing_value, "pt"):
                line_spacing = line_spacing_value.pt / 12.0
            else:
                line_spacing = float(line_spacing_value)

        space_before = _length_to_pt(_first_not_none(*[
            _paragraph_format_value(fmt, "space_before") for fmt in formats
        ]))
        space_after = _length_to_pt(_first_not_none(*[
            _paragraph_format_value(fmt, "space_after") for fmt in formats
        ]))
        alignment = _alignment_to_str(_first_not_none(
            paragraph.alignment,
            *[_paragraph_format_value(fmt, "alignment") for fmt in formats]
        ))
        style_name = paragraph.style.name if paragraph.style else None

        p = Paragraph(
            index=idx,
            text=paragraph.text,
            spans=_extract_spans(paragraph),
            style_name=style_name,
            alignment=alignment,
            first_line_indent_pt=first_indent,
            line_spacing=line_spacing,
            space_before_pt=space_before,
            space_after_pt=space_after,
        )

        ppr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore')
        if ppr is not None:
            p.page_break_before = True

        return p

    from docx.oxml.ns import qn

    para_index = 0
    for block in doc.element.body:
        tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag

        if tag == 'p':
            from docx.text.paragraph import Paragraph as DocxParagraph
            try:
                para = DocxParagraph(block, doc)
                doc_info.paragraphs.append(_extract_paragraph(para, para_index))
            finally:
                para_index += 1

        elif tag == 'tbl':
            rows = block.findall('.//' + qn('w:tr'))
            table_cells = []
            headers = []
            for i, row in enumerate(rows):
                cells = row.findall('.//' + qn('w:tc'))
                row_texts = []
                for cell in cells:
                    cell_text = ''.join(
                        p.text or ''
                        for p in cell.findall('.//' + qn('w:t'))
                    )
                    row_texts.append(cell_text.strip())
                if i == 0:
                    headers = row_texts
                table_cells.append(row_texts)

            if table_cells:
                doc_info.tables.append(TableInfo(
                    index=len(doc_info.tables),
                    rows=len(table_cells),
                    cols=len(table_cells[0]) if table_cells else 0,
                    headers=headers,
                    cells=table_cells,
                ))

    doc_info.total_pages_estimate = max(1, len(doc_info.paragraphs) // 30)
    return doc_info
