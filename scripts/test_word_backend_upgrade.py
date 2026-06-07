import os
import tempfile
import unittest
import warnings
import zipfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
import sys
sys.path.insert(0, str(SCRIPT_DIR))


class WordBackendUpgradeTests(unittest.TestCase):
    def test_parsed_document_has_backend_metadata(self):
        from parsers.docx_parser import ParsedDocument

        doc = ParsedDocument(file_path="sample.docx", file_type="docx")

        self.assertEqual(doc.parser_backend, "python-docx")
        self.assertEqual(doc.parser_warnings, [])

    def test_reader_rejects_unknown_backend(self):
        from reader import read_document

        with self.assertRaises(ValueError) as context:
            read_document("missing.docx", backend="unknown")

        self.assertIn("不支持的解析后端", str(context.exception))

    def test_python_docx_parser_uses_inherited_style_formatting(self):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError:
            self.skipTest("python-docx is not installed")

        from parsers.docx_parser import parse_docx

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "styled.docx"
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "SimSun"
            style.font.size = Pt(12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.first_line_indent = Pt(24)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(3)
            paragraph = doc.add_paragraph("继承样式段落")
            paragraph.style = style
            doc.save(path)

            parsed = parse_docx(str(path))

        target = next(p for p in parsed.paragraphs if p.text == "继承样式段落")
        self.assertEqual(parsed.parser_backend, "python-docx")
        self.assertEqual(target.dominant_font_name(), "SimSun")
        self.assertEqual(target.dominant_font_size_pt(), 12)
        self.assertEqual(target.alignment, "center")
        self.assertEqual(target.first_line_indent_pt, 24)
        self.assertEqual(target.space_before_pt, 6)
        self.assertEqual(target.space_after_pt, 3)

    def test_class_generator_renders_standalone_cls_from_spec(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())

        self.assertIn("\\ProvidesClass{school-thesis}", cls)
        self.assertIn("\\LoadClass", cls)
        self.assertIn("ctexrep", cls)
        self.assertIn("\\newcommand{\\ThesisCoverTitle}", cls)
        self.assertIn("\\titleformat{\\chapter}", cls)
        self.assertIn("\\newenvironment{ChineseAbstract}", cls)
        self.assertIn("\\newcommand{\\ChineseKeywords}", cls)
        self.assertIn("\\renewcommand{\\bibfont}", cls)

    def test_class_generator_uses_explicit_center_blocks_for_unnumbered_titles(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())

        self.assertIn("\\begin{center}", cls)
        self.assertIn("\\end{center}", cls)
        self.assertNotIn("\\chapter*{\\centering", cls)

    def test_class_generator_renders_engineering_skeleton_modules(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())
        lines = cls.splitlines()

        self.assertGreaterEqual(len(lines), 300)
        for marker in [
            "% === Parameter Macros ===",
            "% === Font Setup ===",
            "% === Page Layout ===",
            "% === Header and Footer ===",
            "% === Cover ===",
            "% === Abstracts and Keywords ===",
            "% === Contents and Lists ===",
            "% === Chapter and Section Titles ===",
            "% === Captions ===",
            "% === Bibliography ===",
            "% === Appendix and Acknowledgement ===",
            "% === Theorem and Algorithm Environments ===",
            "% === Declaration Placeholders ===",
        ]:
            self.assertIn(marker, cls)

    def test_class_generator_exposes_thesis_parameter_macros(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())

        for macro in [
            "\\newcommand{\\ThesisBodySize}",
            "\\newcommand{\\ThesisBodyLine}",
            "\\newcommand{\\ThesisBodyFont}",
            "\\newcommand{\\ThesisBodyIndent}",
            "\\newcommand{\\ThesisAbstractTitleSize}",
            "\\newcommand{\\ThesisAbstractTitleFont}",
            "\\newcommand{\\ThesisKeywordFont}",
            "\\newcommand{\\ThesisPrintBibliography}",
            "\\newcommand{\\MakeThesisCover}",
            "\\newcommand{\\ThesisContents}",
            "\\newcommand{\\ThesisFigureList}",
            "\\newcommand{\\ThesisTableList}",
        ]:
            self.assertIn(macro, cls)

    def test_class_generator_exposes_profile_rendering_macros(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())

        for macro in [
            "\\newcommand{\\ThesisSchoolName}",
            "\\newcommand{\\ThesisType}",
            "\\newcommand{\\ThesisSetSchoolName}",
            "\\newcommand{\\ThesisSetType}",
            "\\newcommand{\\MakeThesisProfileCover}",
            "\\newcommand{\\ThesisFrontmatterPage}",
            "\\newcommand{\\ThesisDeclarationPage}",
        ]:
            self.assertIn(macro, cls)

    def test_class_generator_includes_universal_thesis_rules(self):
        from class_generator import render_class
        from extractor import ThesisSpec

        cls = render_class(ThesisSpec())

        for required in [
            "\\RequirePackage{indentfirst}",
            "\\RequirePackage{lastpage}",
            "doi=false",
            "url=false",
            "gbstrict=true",
            "gbalign=center",
            "\\newcommand{\\normcite}{\\parencite}",
            "\\AtBeginEnvironment{figure}{\\def\\@floatboxreset{\\centering}}",
            "\\AtBeginEnvironment{table}{\\def\\@floatboxreset{\\centering}}",
            "\\renewcommand{\\theequation}{\\arabic{chapter}-\\arabic{equation}}",
            "\\renewcommand{\\thefigure}{\\arabic{chapter}.\\arabic{figure}}",
            "\\renewcommand{\\thetable}{\\arabic{chapter}.\\arabic{table}}",
            "\\newtheorem{axiom}[theorem]{公理}",
            "\\newtheorem{proposition}[theorem]{命题}",
            "\\newtheorem{assumption}[theorem]{假设}",
            "\\renewcommand{\\proofname}{证明}",
            "\\renewcommand{\\thechapter}{附录\\Alph{chapter}}",
            "\\renewcommand{\\thesection}{\\Alph{chapter}.\\arabic{section}}",
            "\\renewcommand{\\theequation}{\\alph{chapter}-\\arabic{equation}}",
        ]:
            self.assertIn(required, cls)

    def test_school_profile_defaults_are_generic_and_editable(self):
        from profile import SchoolProfile

        profile = SchoolProfile()

        self.assertEqual(profile.school_name, "<<< 学校名称 >>>")
        self.assertEqual(profile.thesis_type, "<<< 论文类型 >>>")
        self.assertEqual(profile.output_identity.class_name, "school-thesis")
        self.assertGreaterEqual(len(profile.cover_fields), 5)
        self.assertEqual(profile.cover_fields[0].label, "题目")
        self.assertEqual(profile.cover_fields[0].placeholder, "<<< 论文题目 >>>")
        self.assertTrue(any(page.key == "cover" for page in profile.frontmatter_pages))

    def test_profile_yaml_round_trips_scalar_fields_and_cover_fields(self):
        from profile import CoverField, SchoolProfile
        from profile_yaml import dumps_profile_yaml, loads_profile_yaml

        profile = SchoolProfile(
            school_name="测试大学",
            thesis_type="本科毕业设计（论文）",
            degree_level="本科",
            document_title="测试题目",
            cover_fields=[CoverField("题目", "<<< 题目 >>>"), CoverField("学号", "<<< 学号 >>>")],
        )

        text = dumps_profile_yaml(profile)
        loaded = loads_profile_yaml(text, SchoolProfile())

        self.assertIn("school_name: 测试大学", text)
        self.assertIn("  - 题目", text)
        self.assertEqual(loaded.school_name, "测试大学")
        self.assertEqual(loaded.thesis_type, "本科毕业设计（论文）")
        self.assertEqual([field.label for field in loaded.cover_fields], ["题目", "学号"])

    def test_profile_yaml_keeps_existing_values_for_missing_or_empty_fields(self):
        from profile import CoverField, SchoolProfile
        from profile_yaml import loads_profile_yaml

        base = SchoolProfile(school_name="自动识别大学", thesis_type="自动识别类型", cover_fields=[CoverField("姓名", "<<< 姓名 >>>")])
        loaded = loads_profile_yaml("school_name: \nthesis_type: 手工类型\n", base)

        self.assertEqual(loaded.school_name, "自动识别大学")
        self.assertEqual(loaded.thesis_type, "手工类型")
        self.assertEqual([field.label for field in loaded.cover_fields], ["姓名"])

    def test_profile_report_includes_summary_placeholders_and_score(self):
        from profile import CoverField, SchoolProfile
        from profile_report import profile_quality_score, render_profile_report

        profile = SchoolProfile(
            thesis_type="本科毕业设计（论文）",
            degree_level="本科",
            cover_fields=[CoverField("题目", "<<< 题目 >>>"), CoverField("学号", "<<< 学号 >>>")],
        )

        report = render_profile_report(profile)

        self.assertEqual(profile_quality_score(profile), 55)
        self.assertIn("# School Profile Report", report)
        self.assertIn("School name: <<< 学校名称 >>>", report)
        self.assertIn("Thesis type: 本科毕业设计（论文）", report)
        self.assertIn("Quality score: 55/100", report)
        self.assertIn("- school_name", report)
        self.assertIn("1. 题目", report)

    def test_profile_extractor_recognizes_cover_fields_from_tables(self):
        from parsers.docx_parser import ParsedDocument, TableInfo
        from profile_extractor import extract_school_profile

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        parsed.tables.append(TableInfo(index=0, rows=3, cols=2, cells=[["题目", ""], ["学生姓名", ""], ["指导教师职称", ""]]))

        profile = extract_school_profile(parsed)

        self.assertEqual([field.label for field in profile.cover_fields], ["题目", "学生姓名", "指导教师职称"])

    def test_profile_extractor_recognizes_school_name_and_thesis_type(self):
        from parsers.docx_parser import FormattedSpan, Paragraph, ParsedDocument
        from profile_extractor import extract_school_profile

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        parsed.paragraphs.extend([
            Paragraph(index=0, text="示例大学", spans=[FormattedSpan(text="示例大学", font_name="黑体", font_size_pt=22)], alignment="center"),
            Paragraph(index=1, text="本科毕业设计（论文）", spans=[FormattedSpan(text="本科毕业设计（论文）", font_name="黑体", font_size_pt=18)], alignment="center"),
        ])

        profile = extract_school_profile(parsed)

        self.assertEqual(profile.school_name, "示例大学")
        self.assertEqual(profile.thesis_type, "本科毕业设计（论文）")
        self.assertEqual(profile.degree_level, "本科")
        self.assertIn("示例大学", profile.output_identity.readme_title)

    def test_profile_extractor_recognizes_cover_fields_in_order(self):
        from parsers.docx_parser import FormattedSpan, Paragraph, ParsedDocument
        from profile_extractor import extract_school_profile

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        labels = ["题目：", "学生姓名：", "学号：", "学院：", "专业：", "指导教师：", "完成日期："]
        for index, label in enumerate(labels):
            parsed.paragraphs.append(Paragraph(index=index, text=label, spans=[FormattedSpan(text=label, font_name="宋体", font_size_pt=12)], alignment="left"))

        profile = extract_school_profile(parsed)

        self.assertEqual([field.label for field in profile.cover_fields], ["题目", "学生姓名", "学号", "学院", "专业", "指导教师", "完成日期"])
        self.assertEqual(profile.cover_fields[1].placeholder, "<<< 学生姓名 >>>")

    def test_profile_extractor_recognizes_declaration_pages_and_signatures(self):
        from parsers.docx_parser import FormattedSpan, Paragraph, ParsedDocument
        from profile_extractor import extract_school_profile

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        parsed.paragraphs.extend([
            Paragraph(index=0, text="原创性声明", spans=[FormattedSpan(text="原创性声明", font_name="黑体", font_size_pt=16)], alignment="center"),
            Paragraph(index=1, text="本人声明所呈交的毕业设计（论文）是本人在导师指导下完成。", spans=[FormattedSpan(text="本人声明所呈交的毕业设计（论文）是本人在导师指导下完成。", font_name="宋体", font_size_pt=12)], alignment="left"),
            Paragraph(index=2, text="作者签名：        日期：", spans=[FormattedSpan(text="作者签名：        日期：", font_name="宋体", font_size_pt=12)], alignment="left"),
        ])

        profile = extract_school_profile(parsed)

        self.assertEqual(len(profile.declaration_pages), 1)
        self.assertEqual(profile.declaration_pages[0].title, "原创性声明")
        self.assertIn("本人声明", profile.declaration_pages[0].body_text_or_placeholder)
        self.assertEqual(profile.declaration_pages[0].signature_fields, ["作者签名", "日期"])

    def test_format_rule_parser_extracts_chinese_title_instruction(self):
        from format_rule_parser import parse_format_rule

        rule = parse_format_rule("三号黑体居中")

        self.assertIsNotNone(rule)
        self.assertEqual(rule.text_format.font_cn, "\\heiti")
        self.assertEqual(rule.text_format.font_size_pt, 16)
        self.assertEqual(rule.text_format.alignment, "center")

    def test_format_rule_parser_extracts_body_instruction(self):
        from format_rule_parser import parse_format_rule

        rule = parse_format_rule("内容采用小四号宋体，首行缩进两个字，1.5倍行距")

        self.assertIsNotNone(rule)
        self.assertEqual(rule.text_format.font_cn, "\\songti")
        self.assertEqual(rule.text_format.font_size_pt, 12)
        self.assertEqual(rule.text_format.first_line_indent, "2\\ccwd")
        self.assertEqual(rule.text_format.alignment, "justify")

    def test_format_rule_parser_detects_semantic_target(self):
        from format_rule_parser import parse_format_rule

        rule = parse_format_rule("关键词：(小四号、黑体、顶格)")

        self.assertIsNotNone(rule)
        self.assertEqual(rule.semantic_label, "中文关键词标签")
        self.assertEqual(rule.text_format.font_cn, "\\heiti")
        self.assertEqual(rule.text_format.font_size_pt, 12)
        self.assertEqual(rule.text_format.alignment, "left")

    def test_format_rule_parser_does_not_match_size_name_inside_common_words(self):
        from format_rule_parser import parse_format_rule

        rule = parse_format_rule("（点线大小一致且行间距间隔统一，页码字体统一小四）")

        self.assertIsNotNone(rule)
        self.assertEqual(rule.text_format.font_size_pt, 12)

    def test_extract_specs_uses_format_instruction_rule_for_target_style(self):
        from cluster import cluster_paragraphs
        from extractor import extract_specs_from_clusters
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [Paragraph(
            index=0,
            text="摘要（三号黑体居中）",
            spans=[FormattedSpan(text="摘要（三号黑体居中）", font_name="宋体", font_size_pt=12)],
            alignment="left",
        )]
        result = cluster_paragraphs(paragraphs, n_clusters=1)
        spec = extract_specs_from_clusters(result, {result.clusters[0].cluster_id: "中文摘要标题"})

        self.assertEqual(spec.abstract_cn_title.font_cn, "\\heiti")
        self.assertEqual(spec.abstract_cn_title.font_size_pt, 16)
        self.assertEqual(spec.abstract_cn_title.alignment, "center")

    def test_extract_specs_prefers_rule_matching_requested_label(self):
        from cluster import cluster_paragraphs
        from extractor import extract_specs_from_clusters
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(index=0, text="（内容与关键词空一行）", spans=[FormattedSpan(text="（内容与关键词空一行）", font_name="宋体", font_size_pt=14)], alignment="justify"),
            Paragraph(index=1, text="关键词：(小四号、黑体、顶格)", spans=[FormattedSpan(text="关键词：(小四号、黑体、顶格)", font_name="宋体", font_size_pt=12)], alignment="justify"),
        ]
        result = cluster_paragraphs(paragraphs, n_clusters=1)
        spec = extract_specs_from_clusters(result, {result.clusters[0].cluster_id: "中文关键词标签"})

        self.assertEqual(spec.abstract_cn_keyword_label.font_cn, "\\heiti")
        self.assertEqual(spec.abstract_cn_keyword_label.font_size_pt, 12)
        self.assertEqual(spec.abstract_cn_keyword_label.alignment, "left")

    def test_extract_specs_prefers_semantic_rule_over_generic_rule(self):
        from cluster import cluster_paragraphs
        from extractor import extract_specs_from_clusters
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(index=0, text="三号宋体居中", spans=[FormattedSpan(text="三号宋体居中", font_name="宋体", font_size_pt=12)], alignment="center"),
            Paragraph(index=1, text="Abstract（三号加粗）", spans=[FormattedSpan(text="Abstract（三号加粗）", font_name="宋体", font_size_pt=12)], alignment="center"),
        ]
        result = cluster_paragraphs(paragraphs, n_clusters=1)
        spec = extract_specs_from_clusters(result, {result.clusters[0].cluster_id: "英文摘要标题"})

        self.assertEqual(spec.abstract_en_title.font_en, "\\rmfamily")
        self.assertTrue(spec.abstract_en_title.bold)
        self.assertEqual(spec.abstract_en_title.font_size_pt, 16)

    def test_standalone_label_map_uses_semantic_label_from_format_instruction(self):
        from build_standalone import build_label_map
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [Paragraph(
            index=0,
            text="摘要（三号黑体居中）",
            spans=[FormattedSpan(text="摘要（三号黑体居中）", font_name="宋体", font_size_pt=12)],
            alignment="left",
        )]
        result = cluster_paragraphs(paragraphs, n_clusters=1)
        label_map = build_label_map(result)

        self.assertEqual(label_map[result.clusters[0].cluster_id], "中文摘要标题")

    def test_build_standalone_apply_profile_overrides_uses_yaml_then_cli_priority(self):
        from profile import SchoolProfile
        from build_standalone import apply_profile_overrides

        class Args:
            profile = None
            school_name = "CLI大学"
            thesis_type = "CLI论文类型"

        profile = apply_profile_overrides(SchoolProfile(school_name="自动大学", thesis_type="自动类型"), Args())

        self.assertEqual(profile.school_name, "CLI大学")
        self.assertEqual(profile.thesis_type, "CLI论文类型")

    def test_build_standalone_parse_args_supports_profile_and_cli_identity_overrides(self):
        import build_standalone

        parser = build_standalone.create_parser()
        args = parser.parse_args(["input.docx", "--profile", "profile.yaml", "--school-name", "测试大学", "--thesis-type", "测试论文"])

        self.assertEqual(args.profile, "profile.yaml")
        self.assertEqual(args.school_name, "测试大学")
        self.assertEqual(args.thesis_type, "测试论文")

    def test_skill_doc_mentions_profile_report_yaml_and_cli_overrides(self):
        skill_path = SCRIPT_DIR.parent / "SKILL.md"
        text = skill_path.read_text(encoding="utf-8")

        self.assertIn("profile_report.md", text)
        self.assertIn("profile.yaml", text)
        self.assertIn("--school-name", text)
        self.assertIn("--thesis-type", text)

    def test_skill_doc_mentions_profile_driven_school_specific_output(self):
        skill_path = SCRIPT_DIR.parent / "SKILL.md"
        text = skill_path.read_text(encoding="utf-8")

        self.assertIn("SchoolProfile", text)
        self.assertIn("学校专属", text)
        self.assertIn("不硬编码", text)

    def test_build_standalone_imports_profile_extractor_and_project_supports_profile_argument(self):
        import inspect
        import build_standalone
        from standalone_project import generate_standalone_project

        source = inspect.getsource(build_standalone)
        signature = inspect.signature(generate_standalone_project)

        self.assertIn("from profile_extractor import extract_school_profile", source)
        self.assertIn("profile=profile", source)
        self.assertIn("profile", signature.parameters)

    def test_standalone_project_generates_overleaf_zip_structure(self):
        from extractor import ThesisSpec
        from standalone_project import generate_standalone_project
        from packager import package_zip

        with tempfile.TemporaryDirectory() as tmp:
            project_dir = Path(tmp) / "project"
            zip_path = Path(tmp) / "standalone.zip"
            generate_standalone_project(ThesisSpec(), str(project_dir), chapter_count=3)
            package_zip(str(project_dir), str(zip_path))

            with zipfile.ZipFile(zip_path) as archive:
                names = set(archive.namelist())

        self.assertIn("main.tex", names)
        self.assertIn("school-thesis.cls", names)
        self.assertIn("latexmkrc", names)
        self.assertIn("README.md", names)
        self.assertIn("references.bib", names)
        self.assertIn("profile.yaml", names)
        self.assertIn("profile_report.md", names)
        self.assertIn("frontmatter/cover.tex", names)
        self.assertIn("chapters/abstract_cn.tex", names)
        self.assertIn("chapters/abstract_en.tex", names)
        self.assertIn("chapters/chapter1.tex", names)
        self.assertIn("chapters/chapter2.tex", names)
        self.assertIn("chapters/chapter3.tex", names)
        self.assertIn("chapters/conclusion.tex", names)
        self.assertIn("chapters/acknowledgement.tex", names)
        self.assertIn("chapters/references.tex", names)
        self.assertIn("chapters/appendix.tex", names)
        self.assertIn("figures/.keep", names)

    def test_standalone_project_generates_profile_frontmatter_files(self):
        from extractor import ThesisSpec
        from profile import CoverField, DeclarationPage, FrontmatterPage, SchoolProfile
        from standalone_project import generate_standalone_project

        profile = SchoolProfile(
            school_name="示例大学",
            thesis_type="本科毕业设计（论文）",
            cover_fields=[
                CoverField("题目", "<<< 论文题目 >>>"),
                CoverField("学生姓名", "<<< 学生姓名 >>>"),
                CoverField("学号", "<<< 学号 >>>"),
            ],
            frontmatter_pages=[
                FrontmatterPage("cover", "封面", "", False, 0),
                FrontmatterPage("task_book", "任务书", "<<< 请输入任务书内容 >>>", False, 10),
            ],
            declaration_pages=[
                DeclarationPage("declaration_1", "原创性声明", "本人声明测试文本。", ["作者签名", "日期"], True, 100),
            ],
        )

        with tempfile.TemporaryDirectory() as tmp:
            project_dir = Path(tmp) / "project"
            generate_standalone_project(ThesisSpec(), str(project_dir), chapter_count=1, profile=profile)

            main_tex = (project_dir / "main.tex").read_text(encoding="utf-8")
            cover_tex = (project_dir / "frontmatter" / "cover.tex").read_text(encoding="utf-8")
            task_tex = (project_dir / "frontmatter" / "task_book.tex").read_text(encoding="utf-8")
            declaration_tex = (project_dir / "frontmatter" / "declaration_1.tex").read_text(encoding="utf-8")
            readme = (project_dir / "README.md").read_text(encoding="utf-8")
            profile_yaml = (project_dir / "profile.yaml").read_text(encoding="utf-8")
            profile_report = (project_dir / "profile_report.md").read_text(encoding="utf-8")

        self.assertIn("\\input{frontmatter/cover}", main_tex)
        self.assertIn("\\input{frontmatter/task_book}", main_tex)
        self.assertIn("\\input{frontmatter/declaration_1}", main_tex)
        self.assertIn("\\ThesisSetSchoolName{示例大学}", main_tex)
        self.assertIn("\\ThesisSetType{本科毕业设计（论文）}", main_tex)
        self.assertIn("\\ThesisCoverField{学生姓名}{<<< 学生姓名 >>>}", cover_tex)
        self.assertIn("任务书", task_tex)
        self.assertIn("本人声明测试文本。", declaration_tex)
        self.assertIn("作者签名", declaration_tex)
        self.assertIn("示例大学 Thesis Template", readme)
        self.assertIn("school_name: 示例大学", profile_yaml)
        self.assertIn("cover_fields:", profile_yaml)
        self.assertIn("# School Profile Report", profile_report)
        self.assertIn("Quality score:", profile_report)

    def test_auto_cluster_count_uses_thesis_friendly_floor(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = []
        for i in range(146):
            paragraphs.append(Paragraph(
                index=i,
                text=f"段落{i}",
                spans=[FormattedSpan(text=f"段落{i}", font_name="宋体", font_size_pt=12)],
                alignment="justify" if i % 2 else "center",
                first_line_indent_pt=24 if i % 3 == 0 else None,
                line_spacing=1.5 if i % 5 == 0 else 1.0,
            ))

        result = cluster_paragraphs(paragraphs)

        self.assertGreaterEqual(result.n_clusters, 8)

    def test_auto_cluster_count_does_not_emit_convergence_warning_for_duplicate_formats(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(
                index=i,
                text=f"段落{i}",
                spans=[FormattedSpan(text=f"段落{i}", font_name="宋体", font_size_pt=12)],
                alignment="center",
            )
            for i in range(30)
        ]

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            cluster_paragraphs(paragraphs)

        self.assertFalse(any("Number of distinct clusters" in str(item.message) for item in caught))

    def test_format_instruction_cluster_gets_format_label_suggestion(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(index=0, text="三号黑体居中", spans=[FormattedSpan(text="三号黑体居中", font_name="宋体", font_size_pt=12)], alignment="center"),
            Paragraph(index=1, text="内容采用小四号宋体", spans=[FormattedSpan(text="内容采用小四号宋体", font_name="宋体", font_size_pt=12)], alignment="justify"),
            Paragraph(index=2, text="首行缩进两个字", spans=[FormattedSpan(text="首行缩进两个字", font_name="宋体", font_size_pt=12)], alignment="justify"),
        ]

        result = cluster_paragraphs(paragraphs, n_clusters=1)

        self.assertEqual(result.clusters[0].suggested_label, "格式说明")
        self.assertGreaterEqual(result.clusters[0].suggested_label_confidence, 0.8)
        self.assertGreater(result.clusters[0].classification_counts.get("format_annotation", 0), 0)

    def test_body_cluster_gets_body_label_suggestion(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(
                index=i,
                text=f"这是正文内容{i}",
                spans=[FormattedSpan(text=f"这是正文内容{i}", font_name="宋体", font_size_pt=12)],
                alignment="justify",
                first_line_indent_pt=24,
                line_spacing=1.5,
            )
            for i in range(6)
        ]

        result = cluster_paragraphs(paragraphs, n_clusters=1)

        self.assertEqual(result.clusters[0].suggested_label, "正文段落")
        self.assertGreaterEqual(result.clusters[0].suggested_label_confidence, 0.75)

    def test_title_cluster_gets_title_label_suggestion(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [Paragraph(
            index=0,
            text="本科毕业设计（论文）",
            spans=[FormattedSpan(text="本科毕业设计（论文）", font_name="黑体", font_size_pt=26, bold=True)],
            alignment="center",
            line_spacing=2,
        )]

        result = cluster_paragraphs(paragraphs, n_clusters=1)

        self.assertEqual(result.clusters[0].suggested_label, "论文大标题")
        self.assertGreaterEqual(result.clusters[0].suggested_label_confidence, 0.9)

    def test_report_includes_classification_summary_and_suggested_label_map(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph, ParsedDocument
        from reporter import write_report

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        parsed.paragraphs.extend([
            Paragraph(index=0, text="本科毕业设计（论文）", spans=[FormattedSpan(text="本科毕业设计（论文）", font_name="黑体", font_size_pt=26, bold=True)], alignment="center"),
            Paragraph(index=1, text="三号黑体居中", spans=[FormattedSpan(text="三号黑体居中", font_name="宋体", font_size_pt=12)], alignment="center"),
        ])
        result = cluster_paragraphs(parsed.paragraphs, n_clusters=2)

        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "report.md"
            write_report(parsed, result, str(report))
            text = report.read_text(encoding="utf-8")

        self.assertIn("## Classification Summary", text)
        self.assertIn("## Suggested Label Map", text)
        self.assertIn("suggested_label", text)

    def test_word_com_parser_treats_undefined_font_size_as_missing(self):
        from parsers.word_com_parser import _font_size

        class Font:
            Size = 9999999

        self.assertIsNone(_font_size(Font()))

    def test_word_com_parser_availability_returns_bool(self):
        from parsers.word_com_parser import is_word_com_available

        self.assertIsInstance(is_word_com_available(), bool)

    def test_cluster_profiles_include_format_signature(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph

        paragraphs = [
            Paragraph(
                index=0,
                text="正文段落",
                spans=[FormattedSpan(text="正文段落", font_name="宋体", font_size_pt=12)],
                style_name="Normal",
                alignment="left",
                first_line_indent_pt=24,
                line_spacing=1.5,
                space_before_pt=0,
                space_after_pt=0,
            )
        ]

        result = cluster_paragraphs(paragraphs, n_clusters=1)

        self.assertIn("style=Normal", result.clusters[0].format_signature)
        self.assertIn("font=宋体", result.clusters[0].format_signature)

    def test_reporter_writes_backend_and_cluster_report(self):
        from cluster import cluster_paragraphs
        from parsers.docx_parser import FormattedSpan, Paragraph, ParsedDocument
        from reporter import write_report

        parsed = ParsedDocument(file_path="sample.docx", file_type="docx")
        parsed.paragraphs.append(Paragraph(
            index=0,
            text="正文段落",
            spans=[FormattedSpan(text="正文段落", font_name="宋体", font_size_pt=12)],
            style_name="Normal",
            alignment="left",
        ))
        result = cluster_paragraphs(parsed.paragraphs, n_clusters=1)

        with tempfile.TemporaryDirectory() as tmp:
            report = Path(tmp) / "report.md"
            write_report(parsed, result, str(report))
            text = report.read_text(encoding="utf-8")

        self.assertIn("Parser backend", text)
        self.assertIn("python-docx", text)
        self.assertIn("Format Clusters", text)

    def test_reference_project_preserves_example_architecture_and_removes_build_artifacts(self):
        from reference_project import build_reference_project

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            reference = tmp_path / "example-reference.zip"
            with zipfile.ZipFile(reference, "w") as archive:
                archive.writestr("ExampleUniversity-Thesis.cls", "% class")
                archive.writestr("main.tex", "\\documentclass{ExampleUniversity-Thesis}\\begin{document}x\\end{document}")
                archive.writestr("latexmkrc", "$xelatex='xelatex';")
                archive.writestr("Chapters/Abstract.tex", "abstract")
                archive.writestr("Fonts/fzsong.ttf", "font")
                archive.writestr("Pictures/University-logo.pdf", "pdf-source")
                archive.writestr("References/References.bib", "@book{x}")
                archive.writestr("Build/main.pdf", "compiled")
                archive.writestr("main.aux", "aux")
                archive.writestr("main.log", "log")
                archive.writestr("main.pdf", "compiled")
            report = tmp_path / "report.md"
            report.write_text("# recognition", encoding="utf-8")

            zip_path = build_reference_project(str(reference), str(tmp_path / "out"), str(report))

            with zipfile.ZipFile(zip_path) as archive:
                names = set(archive.namelist())

        self.assertIn("ExampleUniversity-Thesis.cls", names)
        self.assertIn("main.tex", names)
        self.assertIn("latexmkrc", names)
        self.assertIn("Chapters/Abstract.tex", names)
        self.assertIn("Fonts/fzsong.ttf", names)
        self.assertIn("Pictures/University-logo.pdf", names)
        self.assertIn("References/References.bib", names)
        self.assertIn("thesis-template-recognition-report.md", names)
        self.assertIn("THESIS_TEMPLATE_GENERATION.md", names)
        self.assertNotIn("Build/main.pdf", names)
        self.assertNotIn("main.aux", names)
        self.assertNotIn("main.log", names)
        self.assertNotIn("main.pdf", names)

    def test_reference_project_builds_sanitized_zip_with_report(self):
        from reference_project import build_reference_project

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            reference = tmp_path / "reference.zip"
            with zipfile.ZipFile(reference, "w") as archive:
                archive.writestr("main.tex", "\\title{旧题目}\n\\chapter{第一章 原内容}\n")
                archive.writestr("build.aux", "temporary")
                archive.writestr("Fonts/font.txt", "font")
            report = tmp_path / "thesis-template-report.md"
            report.write_text("# report", encoding="utf-8")

            zip_path = build_reference_project(str(reference), str(tmp_path / "out"), str(report))

            self.assertTrue(zip_path.exists())
            with zipfile.ZipFile(zip_path) as archive:
                names = archive.namelist()
                main_tex = archive.read("main.tex").decode("utf-8")

        self.assertIn("main.tex", names)
        self.assertIn("Fonts/font.txt", names)
        self.assertIn("thesis-template-recognition-report.md", names)
        self.assertNotIn("build.aux", names)
        self.assertIn("<<< 论文题目 >>>", main_tex)
        self.assertIn("<<< 章节标题 >>>", main_tex)


if __name__ == "__main__":
    unittest.main()
